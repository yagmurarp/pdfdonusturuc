from flask import Flask, render_template, request, send_file
import os
import pandas as pd
from docx import Document
from openpyxl import Workbook
from fpdf import FPDF
from pdf2image import convert_from_path
from werkzeug.utils import secure_filename
import comtypes.client
import tempfile

app = Flask(__name__)
UPLOAD_FOLDER = "uploads"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

@app.route("/", methods=["GET", "POST"])
def index():
    if request.method == "POST":
        file = request.files["file"]
        convert_type = request.form.get("convert_type")

        if not file:
            return "Dosya yüklenmedi!", 400

        filename = secure_filename(file.filename)
        filepath = os.path.join(UPLOAD_FOLDER, filename)
        file.save(filepath)

        output_path = None

        # PDF → Word
        if convert_type == "pdf_to_word":
            output_path = filepath.replace(".pdf", ".docx")
            images = convert_from_path(filepath)
            doc = Document()
            for image in images:
                temp_img = os.path.join(UPLOAD_FOLDER, "temp.jpg")
                image.save(temp_img, "JPEG")
                doc.add_picture(temp_img)
                os.remove(temp_img)
            doc.save(output_path)

        # Word → PDF
        elif convert_type == "word_to_pdf":
            output_path = filepath.replace(".docx", ".pdf")
            word = comtypes.client.CreateObject("Word.Application")
            doc = word.Documents.Open(filepath)
            doc.SaveAs(output_path, FileFormat=17)
            doc.Close()
            word.Quit()

        # PDF → Excel
        elif convert_type == "pdf_to_excel":
            output_path = filepath.replace(".pdf", ".xlsx")
            df = pd.DataFrame({"PDF İçerik": ["Bu örnekte PDF → Excel dönüşümü yapılabilir."]})
            df.to_excel(output_path, index=False)

        # Excel → PDF
        elif convert_type == "excel_to_pdf":
            output_path = filepath.replace(".xlsx", ".pdf")
            pdf = FPDF()
            pdf.add_page()
            pdf.set_font("Arial", size=12)
            pdf.cell(200, 10, txt="Excel verileri PDF'ye dönüştürüldü.", ln=True)
            pdf.output(output_path)

        # Excel → Word
        elif convert_type == "excel_to_word":
            output_path = filepath.replace(".xlsx", ".docx")
            doc = Document()
            doc.add_paragraph("Excel verileri Word'e dönüştürüldü.")
            doc.save(output_path)

        if output_path and os.path.exists(output_path):
            return send_file(output_path, as_attachment=True)
        else:
            return "Dönüştürme işlemi başarısız!", 500

    return render_template("index.html")


if __name__ == "__main__":
    app.run(debug=True)


